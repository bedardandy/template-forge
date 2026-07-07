#!/usr/bin/env python3
"""Export an assembled document to a clean .docx, driven by the DECLARED
structured-segment layer (engine/segments.py) — not by re-parsing Markdown
punctuation.

    python3 engine/export_docx.py --manifest manifests/real_estate__deed.json \
            --slots fixtures/deed_quitclaim.json --out /tmp/deed.docx

Each segment carries a structural `role` (title / heading / body / recital /
numbered_item / signature_block / notary_block / divider / chrome). This module
owns the role -> presentation mapping, so the blocks stay semantic and only the
exporter decides how a role looks. An HTML or document-model exporter would map
the same roles to its own tags.

FORMATTING MODEL (calibrated to the firm's real .docx corpus, 2026-06-20):
- Times New Roman 12, justified body, centered title/caption — the corpus norm.
- Inter-paragraph separation is carried by paragraph spacing and (for court
  filings) line spacing, NEVER by stacked empty paragraphs or 3+-space padding.
  * court filings (pleading/motion/divorce complaint, and litigation/criminal/
    family affidavits) are double-spaced with no space-after — the court norm;
  * everything else is single-spaced with a small space-after between paragraphs.
- Section TITLES/HEADINGS carry keep-with-next so a heading never strands at the
  bottom of a page away from the first paragraph of its section.
- Signature/notary/caption blocks carry keep-together (and table rows can't
  split) so an execution block is never broken across a page.
- Column alignment uses TAB STOPS, not runs of spaces: any 3+-space run in the
  authored text is converted to a real tab against declared tab stops, and
  captions render as borderless 2-column tables.

This is a FLAT delivery docx. The reversible-template target -- DOCX with
<w:sdt> content controls per typed slot -- is a separate, heavier build (see
project_docujsonl_templating).
"""
import argparse
import json
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import assemble as A
import paths as _paths
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BOLD = re.compile(r"(\*\*.*?\*\*)")
MULTISPACE = re.compile(r" {3,}")          # 3+ spaces == column padding -> a tab
                                           # (2 spaces == sentence spacing, kept)

ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

# Tab stops used when a line carries tab-aligned columns (e.g. a witness line
# "____  Address: ____" or a two-column signature). Points from the left margin.
TAB_STOPS = [(Inches(3.25), WD_TAB_ALIGNMENT.LEFT),
             (Inches(5.0), WD_TAB_ALIGNMENT.LEFT),
             (Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)]

# Line spacing: SINGLE is the default (the corpus is ~50/50 on pleadings, the
# clean exemplars are single, and single reads cleanest). A filing that a court
# requires to be double-spaced opts in with `"double_spaced": true` in its slots.
def is_letter(document_type: str | None) -> bool:
    """Correspondence (demand / engagement / cover / discovery letters). A real
    letter has NO centered document title — it opens with the date/letterhead —
    so the exporter drops the title segment for these."""
    return bool(document_type) and document_type.endswith("_letter")


# role -> how it looks. `heading` = Word heading level (a real Word header, so it
# shows in the navigation pane and carries keep-with-next); else a Normal
# paragraph. Indents/space are in points. `per_line` splits a multi-line segment
# (signature/notary geometry) into one paragraph per line and keeps the lines
# together. `keep_next` keeps the paragraph with the following one. `keep_lines`
# keeps the paragraph's own lines on a single page.
ROLE_PRESENTATION = {
    "title":           {"heading": 1, "align": "center", "size": 13, "keep_next": True},
    "heading":         {"heading": 2, "align": "left", "size": 12, "keep_next": True},
    "body":            {"align": "justify", "gap": True},
    "recital":         {"align": "justify", "left_indent": 24, "gap": True},
    "numbered_item":   {"align": "justify", "left_indent": 36, "first_line_indent": -18, "gap": True},
    "signature_block": {"align": "left", "per_line": True, "keep_lines": True, "lead": 10},
    "notary_block":    {"align": "left", "per_line": True, "keep_lines": True, "lead": 10},
    "divider":         {"align": "center"},
    "chrome":          {"align": "left", "per_line": True, "keep_lines": True},
}
DEFAULT_PRES = {"align": "left", "gap": True}


def tabify(text: str) -> str:
    return MULTISPACE.sub("\t", text)


def add_runs(p, text):
    """Add runs to a paragraph, honoring **bold** and converting 3+-space column
    padding into real tabs (emitted as <w:tab/> runs, not space text)."""
    text = tabify(text)
    for i, piece in enumerate(text.split("\t")):
        if i:
            p.add_run().add_tab()
        for seg in BOLD.split(piece):
            if not seg:
                continue
            if seg.startswith("**") and seg.endswith("**"):
                p.add_run(seg[2:-2]).bold = True
            else:
                p.add_run(seg)


def _apply_tabstops(p):
    for pos, al in TAB_STOPS:
        p.paragraph_format.tab_stops.add_tab_stop(pos, al)


def style_heading(h, size, gap):
    """Recorded instruments are black Times New Roman, not Word's theme-blue
    Calibri headings; headings stay with their first paragraph and are tightly
    spaced (the Word heading styles' large space-before is overridden)."""
    for r in h.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.bold = True
    pf = h.paragraph_format
    pf.keep_with_next = True
    pf.line_spacing = 1.0
    pf.space_before = Pt(0 if h._p.getprevious() is None else (8 if gap else 4))
    pf.space_after = Pt(gap if gap else 2)


def _emit(doc, role, line, pres, align, gap, first=True, last=True):
    if pres.get("heading"):
        h = doc.add_heading("", level=pres["heading"])
        add_runs(h, line)
        h.alignment = ALIGN[align]
        # title gets a clear gap after it; section headings a small one
        style_heading(h, pres.get("size", 12), 12 if pres["heading"] == 1 else 4)
        if "\t" in tabify(line):
            _apply_tabstops(h)
        return h
    p = doc.add_paragraph()
    p.alignment = ALIGN[align]
    pf = p.paragraph_format
    if pres.get("left_indent"):
        pf.left_indent = Pt(pres["left_indent"])
    if pres.get("first_line_indent"):
        pf.first_line_indent = Pt(pres["first_line_indent"])
    if pres.get("per_line"):
        # signature/notary geometry: lines stay together; the block is led in
        # with a small space-before (not a blank paragraph) and separated after.
        pf.keep_together = True
        if first and pres.get("lead"):
            pf.space_before = Pt(pres["lead"])
        if not last:
            pf.keep_with_next = True
        pf.space_after = Pt(gap if (last and gap) else 0)
    else:
        if pres.get("keep_lines"):
            pf.keep_together = True
        pf.space_after = Pt(gap if (gap and pres.get("gap")) else 0)
    add_runs(p, line)
    if "\t" in tabify(line):
        _apply_tabstops(p)
    return p


def _row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:cantSplit")) is None:
        trPr.append(OxmlElement("w:cantSplit"))


def _emit_table(doc, rows, borderless=True):
    rows = [r for r in rows if any(c.strip() for c in r)]
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncols)
    if not borderless:
        t.style = "Table Grid"
    t.autofit = True
    for ri, r in enumerate(rows):
        _row_cant_split(t.rows[ri])          # caption never splits across a page
        for ci in range(ncols):
            cell = t.cell(ri, ci)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            for k, ln in enumerate((r[ci] if ci < len(r) else "").split("\n")):
                if k:
                    p.add_run().add_break()
                add_runs(p, ln)


def _set_default_font(doc, name="Times New Roman", pt=12):
    """Make Times New Roman the document's *default* font (docDefaults), not the
    theme font, so every run resolves to TNR even if it carries no explicit or
    style font — the document keeps one consistent font throughout."""
    el = doc.styles.element
    dd = el.find(qn("w:docDefaults")) or el.makeelement(qn("w:docDefaults"), {})
    if dd.getparent() is None:
        el.insert(0, dd)
    rprd = dd.find(qn("w:rPrDefault"))
    if rprd is None:
        rprd = OxmlElement("w:rPrDefault"); dd.append(rprd)
    rpr = rprd.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr"); rprd.append(rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.insert(0, rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), name)
    for a in ("w:asciiTheme", "w:hAnsiTheme", "w:cstheme", "w:eastAsiaTheme"):
        if rfonts.get(qn(a)) is not None:
            del rfonts.attrib[qn(a)]
    sz = rpr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz"); rpr.append(sz)
    sz.set(qn("w:val"), str(pt * 2))


def segments_to_docx(segments, out, double_spaced=False, drop_title=False):
    doc = Document()
    _set_default_font(doc)
    n = doc.styles["Normal"]
    n.font.name = "Times New Roman"
    n.font.size = Pt(12)
    npf = n.paragraph_format
    npf.space_after = Pt(0)
    npf.space_before = Pt(0)
    npf.line_spacing = 2.0 if double_spaced else 1.0
    # in a double-spaced doc the line spacing carries the gap, so no space-after;
    # in a single-spaced doc paragraphs are separated by an 8pt space-after.
    gap = 0 if double_spaced else 8
    for s in segments:
        role = s.get("role", "body")
        if drop_title and role == "title":     # letters open with the date, no title
            continue
        if s.get("rows") is not None:          # caption / table primitive
            _emit_table(doc, s["rows"], borderless=(role == "caption"))
            continue
        pres = ROLE_PRESENTATION.get(role, DEFAULT_PRES)
        align = s.get("align") or pres.get("align", "left")
        text = s.get("text", "")
        if pres.get("per_line"):
            lines = [ln for ln in text.split("\n") if ln.strip()]
            for i, line in enumerate(lines):
                _emit(doc, role, line, pres, align, gap,
                      first=(i == 0), last=(i == len(lines) - 1))
        else:
            # one Word paragraph; internal soft newlines collapse to spaces
            _emit(doc, role, " ".join(text.split("\n")), pres, align, gap)
    doc.save(out)
    return out


def export(manifest_path: Path, slots: dict, out) -> str:
    """Assemble and write a .docx from the declared segment layer. Single-spaced
    by default; pass `"double_spaced": true` in slots for a filing that must be
    double-spaced. Letters drop the centered title."""
    _, record = A.assemble_record(Path(manifest_path), slots, trail=False)
    double = bool(slots.get("double_spaced"))
    drop_title = is_letter(record.get("document_type"))
    return segments_to_docx(record["segments"], out,
                            double_spaced=double, drop_title=drop_title)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--pack", help="template pack directory "
                    "(default: $TEMPLATE_FORGE_PACK or the bundled example_pack)")
    ap.add_argument("--manifest", required=True)
    ap.add_argument("--slots", required=True)
    ap.add_argument("--out", required=True)
    a = ap.parse_args()
    if a.pack:
        _paths.set_pack(a.pack)
    out = export(Path(a.manifest), json.load(open(a.slots)), a.out)
    print(f"wrote {out}")


if __name__ == "__main__":
    main()
