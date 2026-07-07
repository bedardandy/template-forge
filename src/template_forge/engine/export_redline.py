#!/usr/bin/env python3
"""Export a GRANULAR, EXPLAINED redline .docx: the firm's blank-form base marked
up into the filled instance with Word tracked changes AT THE WORD/RUN LEVEL (not
whole-paragraph replacement), plus Word COMMENTS that explain the DECISION behind
every input and every intentional omission.

    python3 engine/export_redline.py --manifest manifests/real_estate__trustee_deed.json \
            --slots fixtures/deed_trustee.json --out /tmp/trustee.redline.docx

WHAT IT PRODUCES (a review/diligence artifact, not the delivery doc):
- BASE   = the same instrument with every fill-in shown as its blank-form marker
           `[[ slot ]]` (assembled through the identical engine path, so gating,
           auto-numbering and captions match the filled doc one-for-one).
- REDLINE = a word-level diff base -> filled: each filled value is a tracked
           INSERTION (w:ins) replacing the struck marker (w:del). Granular, so a
           single changed name inside a recital shows as one insertion, not a
           rewritten paragraph.
- INPUT COMMENTS  = on each inserted value, a Word comment drawn from the decision
           record's fact ledger: the value, WHERE it came from (attorney fact /
           template pin / engine-derived / template default), and the clause's
           statutory authority. This answers "why is this value here?".
- OMISSION COMMENTS = every clause the engine CONDITIONALLY left out (a `when:`
           gate that evaluated false) is shown as struck base text with a comment
           naming the gate, the facts it read, and the rule note. This answers
           "what was deliberately NOT included, and why?".

The base and the filled document come from engine/assemble.assemble_record, so
this exporter never reinterprets clause text — it only diffs and annotates what
the engine already decided. Tracked-change author/date are deterministic by
default (pass --author/--date to override) so rebuilds don't churn.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
import zipfile
from difflib import SequenceMatcher
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import assemble as A  # noqa: E402
import export_docx as X  # noqa: E402  (reuse font + role presentation)
import render as R  # noqa: E402
import segments as S  # noqa: E402
from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Pt, RGBColor  # noqa: E402

DEF_AUTHOR = "template-forge"
DEF_DATE = "2026-06-28T00:00:00Z"

ALIGN = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
         "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}
MARKER = re.compile(r"\[\[\s*(\w+)\s*\]\]")
# the date slots feed derive_flags(post_2019); keep them real in the blank base so
# blanking them can't silently flip a statutory variant's gating.
FLAG_DATE_SLOTS = {"execution_date", "signing_date", "date"}
NUMERICISH = re.compile(r"^[\$\s]*[\d][\d,\.]*\s*%?$")   # a slot a |int/|float would touch


# ---------------------------------------------------------------------------
# context + base/blank build
# ---------------------------------------------------------------------------
_PRINTED = re.compile(r"\{\{\s*([a-zA-Z_]\w*)")   # {{ var }} / {{ var | filter }}
_LITMARK = re.compile(r"\[\[\s*(\w+)\s*\]\]")      # literal [[ marker ]] in source


def _printed_vars(filled_record: dict, manifest: dict, blocks: dict) -> set:
    """Variables the template PRINTS (anything in a `{{ }}` output expression or a
    `[[ marker ]]`), across every block that actually rendered, plus the title.
    This is the set to blank in the base: a printed variable is a fill-in value;
    a variable that appears ONLY in `{% if %}`/`{% for %}` is a gating flag and is
    deliberately NOT collected, so blanking can never change which clauses fire.
    Declared `slots` metadata is curated and misses party names ({{ grantee_name }}
    is referenced but not always declared), so we read the raw templates."""
    names: set = set()
    for d in filled_record.get("decisions", []):
        if not d.get("included"):
            continue
        bid = d.get("block_id")
        b = blocks.get(bid)
        if not b:
            continue
        try:
            raw = R.block_text(b, d.get("text_source", "dominant"))
        except Exception:
            continue
        names |= set(_PRINTED.findall(raw)) | set(_LITMARK.findall(raw))
    title = manifest.get("title") or ""
    names |= set(_PRINTED.findall(title)) | set(_LITMARK.findall(title))
    return names


def _blank_fixture(slots: dict, names: set, keep_numeric: bool) -> dict:
    """Replace each ACTUALLY-FILLED value slot with its blank-form marker. Slots
    the matter never supplied are left absent so they render as `[[ slot ]]` in
    BOTH base and filled (an unfilled marker is not a change) — and a slot used in
    a `{% if slot %}` guard keeps the same truthiness on both sides (filled real
    value and the marker string are both truthy; an absent slot is falsy on both),
    so gating stays identical and every diff maps to a real input."""
    out = dict(slots)
    for n in names:
        if n not in slots or n in FLAG_DATE_SLOTS:
            continue
        if keep_numeric and isinstance(out.get(n), (int, float)):
            continue
        if keep_numeric and isinstance(out.get(n), str) and NUMERICISH.match(out.get(n) or ""):
            continue
        out[n] = f"[[ {n} ]]"
    return out


def build_records(manifest_path: Path, slots: dict):
    """(filled_record, base_record) assembled through the identical engine path."""
    manifest = json.loads(Path(manifest_path).read_text())
    blocks = R.load_blocks()
    _, filled = A.assemble_record(Path(manifest_path), slots, trail=False)
    names = _printed_vars(filled, manifest, blocks)
    # never blank a discriminator: it is never a printed fill-in, and blanking it
    # would change a `{% if disc == 'x' %}` branch and desync base from filled.
    names -= set((manifest.get("discriminators") or {}).keys())
    base = None
    for keep_numeric in (False, True):
        try:
            _, base = A.assemble_record(Path(manifest_path),
                                        _blank_fixture(slots, names, keep_numeric),
                                        trail=False)
            break
        except SystemExit:
            base = None
        except Exception:
            base = None
    return filled, base, manifest, blocks


# ---------------------------------------------------------------------------
# comment composition from the decision record
# ---------------------------------------------------------------------------
_SRC = {"slots_file": "attorney-provided fact",
        "manifest_pin": "pinned by this template (discriminator)",
        "derived": "derived by the engine",
        "manifest_default": "template default",
        "unknown": "unsourced"}


def input_comment(name: str, facts: dict, authority: str | None) -> str:
    e = facts.get(name) or {}
    val = e.get("value")
    src = _SRC.get(e.get("source", "unknown"), e.get("source", "unknown"))
    lines = [f"INPUT — {name} = {val!r}", f"Source: {src}"]
    if e.get("detail"):
        lines.append(f"  {e['detail']}")
    for k in ("basis", "author", "date"):
        if e.get(k):
            lines.append(f"  {k}: {e[k]}")
    if authority:
        lines.append(f"Clause authority: {authority}")
    return "\n".join(lines)


def omission_comment(d: dict) -> str:
    lines = [f"INTENTIONALLY OMITTED — {d.get('function') or d.get('block_id')}"]
    if d.get("note"):
        lines.append(f"Reason: {d['note']}")
    if d.get("when") is not None:
        lines.append(f"Gate: ({d['when']}) -> {d.get('when_result')}")
        fc = d.get("facts_considered") or {}
        if fc:
            lines.append("Facts read: " + ", ".join(f"{k}={v!r}" for k, v in fc.items()))
    if d.get("statutory_citation"):
        lines.append(f"Authority if used: {d['statutory_citation']}")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# OOXML run / tracked-change / comment builders
# ---------------------------------------------------------------------------
def _tok(text: str) -> list:
    """Tokenize to word+trailing-spaces units (tabs/newlines kept separate for
    w:tab/w:br). Trailing spaces are glued to their word so a lone space can't be
    matched as a diff anchor between base and filled — that fragmentation would
    split a multi-word value's `[[ marker ]]` across ops and lose its comment."""
    return re.findall(r"\t|\n|\S+[ ]*|[ ]+", text or "")


def _run(token: str, deleted: bool = False, rpr_extra=None):
    r = OxmlElement("w:r")
    if rpr_extra is not None:
        r.append(rpr_extra)
    if token == "\t":
        r.append(OxmlElement("w:tab"))
    elif token == "\n":
        r.append(OxmlElement("w:br"))
    else:
        t = OxmlElement("w:delText" if deleted else "w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = token
        r.append(t)
    return r


def _rpr(strike=False, color=None, italic=False):
    if not (strike or color or italic):
        return None
    rpr = OxmlElement("w:rPr")
    if italic:
        rpr.append(OxmlElement("w:i"))
    if strike:
        rpr.append(OxmlElement("w:strike"))
    if color:
        c = OxmlElement("w:color"); c.set(qn("w:val"), color); rpr.append(c)
    return rpr


class _Ids:
    def __init__(self):
        self.change = 0
        self.comment = 0
    def next_change(self):
        self.change += 1
        return self.change
    def next_comment(self):
        self.comment += 1
        return self.comment


def _ins(tokens, ids, author, date):
    el = OxmlElement("w:ins")
    el.set(qn("w:id"), str(ids.next_change()))
    el.set(qn("w:author"), author)
    el.set(qn("w:date"), date)
    for tk in tokens:
        el.append(_run(tk))
    return el


def _del(tokens, ids, author, date, strike_style=True):
    el = OxmlElement("w:del")
    el.set(qn("w:id"), str(ids.next_change()))
    el.set(qn("w:author"), author)
    el.set(qn("w:date"), date)
    for tk in tokens:
        el.append(_run(tk, deleted=True))
    return el


def _equal(tokens):
    return [_run(tk) for tk in tokens]


# ---------------------------------------------------------------------------
# paragraph emit
# ---------------------------------------------------------------------------
def _new_par(doc, role):
    pres = X.ROLE_PRESENTATION.get(role, X.DEFAULT_PRES)
    p = doc.add_paragraph()
    p.alignment = ALIGN[pres.get("align", "left")]
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    if pres.get("left_indent"):
        pf.left_indent = Pt(pres["left_indent"])
    if role in ("title", "heading"):
        for run_default in ():
            pass
    return p, pres


def _strip_bold(s: str) -> str:
    return s.replace("**", "")


def emit_paragraph(doc, role, base_text, filled_text, facts, authority, ids,
                   comments, author, date):
    """One redlined paragraph: word-level diff base->filled, comment each filled
    input. Returns the docx paragraph."""
    p, pres = _new_par(doc, role)
    pe = p._p
    bt, ft = _tok(_strip_bold(base_text)), _tok(_strip_bold(filled_text))
    sm = SequenceMatcher(None, bt, ft, autojunk=False)
    bold = role in ("title", "heading")
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            for r in _equal(ft[j1:j2]):
                if bold:
                    r.insert(0, _rpr() or OxmlElement("w:rPr"))
                pe.append(r)
        elif tag == "delete":
            pe.append(_del(bt[i1:i2], ids, author, date))
        elif tag == "insert":
            _emit_ins_with_comment(pe, [], ft[j1:j2], facts, authority, ids,
                                   comments, author, date)
        elif tag == "replace":
            pe.append(_del(bt[i1:i2], ids, author, date))
            _emit_ins_with_comment(pe, bt[i1:i2], ft[j1:j2], facts, authority,
                                   ids, comments, author, date)
    return p


def _emit_ins_with_comment(pe, base_tokens, ins_tokens, facts, authority, ids,
                           comments, author, date):
    """Append an insertion; if the base side it replaced carried `[[ slot ]]`
    markers, anchor a comment (per slot) explaining the input decision."""
    names = MARKER.findall("".join(base_tokens))
    # de-dup, keep order
    seen, order = set(), []
    for n in names:
        if n not in seen:
            seen.add(n); order.append(n)
    cids = []
    for n in order:
        cid = ids.next_comment()
        cids.append(cid)
        comments.append((cid, author, date, input_comment(n, facts, authority)))
        crs = OxmlElement("w:commentRangeStart"); crs.set(qn("w:id"), str(cid))
        pe.append(crs)
    pe.append(_ins(ins_tokens, ids, author, date))
    for cid in cids:
        cre = OxmlElement("w:commentRangeEnd"); cre.set(qn("w:id"), str(cid))
        pe.append(cre)
        ref = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        st = OxmlElement("w:rStyle"); st.set(qn("w:val"), "CommentReference"); rpr.append(st)
        ref.append(rpr)
        cr = OxmlElement("w:commentReference"); cr.set(qn("w:id"), str(cid))
        ref.append(cr)
        pe.append(ref)


def emit_omission(doc, d, ctx, env, blocks, ids, comments, author, date):
    """Render a conditionally-omitted block as struck base text + an omission
    comment naming the gate and facts. Returns True if anything was emitted."""
    bid = d.get("block_id")
    block = blocks.get(bid)
    if not block:
        return False
    try:
        raw = R.block_text(block, "dominant")
        text = R.render_text(env, raw, ctx)
    except Exception:
        return False
    paras = S._split_paragraphs(text) or [text]
    cid = ids.next_comment()
    comments.append((cid, author, date, omission_comment(d)))
    first = True
    for para in paras:
        p, _ = _new_par(doc, (d.get("structure") or {}).get("role") or "body")
        pe = p._p
        if first:
            crs = OxmlElement("w:commentRangeStart"); crs.set(qn("w:id"), str(cid))
            pe.append(crs)
        pe.append(_del(_tok(_strip_bold(para)), ids, author, date))
        if first:
            cre = OxmlElement("w:commentRangeEnd"); cre.set(qn("w:id"), str(cid))
            pe.append(cre)
            ref = OxmlElement("w:r")
            rpr = OxmlElement("w:rPr")
            st = OxmlElement("w:rStyle"); st.set(qn("w:val"), "CommentReference"); rpr.append(st)
            ref.append(rpr)
            cr = OxmlElement("w:commentReference"); cr.set(qn("w:id"), str(cid))
            ref.append(cr)
            pe.append(ref)
        first = False
    return True


# ---------------------------------------------------------------------------
# header / orientation
# ---------------------------------------------------------------------------
def _header(doc, filled, author):
    title = filled.get("title") or filled.get("template_id")
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(f"REDLINE — {title}")
    r.bold = True; r.font.size = Pt(13)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(
        f"{filled.get('practice_area')}/{filled.get('document_type')}  ·  "
        f"template {filled.get('template_id')}")
    rs.italic = True; rs.font.size = Pt(9); rs.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    leg = doc.add_paragraph()
    rl = leg.add_run(
        "Base = blank-form markers [[ slot ]]; insertions = the inputs filled for "
        "this matter (comment on each = the decision behind it); struck text = "
        "clauses the engine intentionally omitted (comment = the gate + facts). "
        f"Tracked-change author: {author}.")
    rl.italic = True; rl.font.size = Pt(9); rl.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph().add_run("").add_break()


# ---------------------------------------------------------------------------
# comments part (zip post-processing)
# ---------------------------------------------------------------------------
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CT_COMMENTS = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
REL_COMMENTS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"


def _comments_xml(comments) -> str:
    from xml.sax.saxutils import escape
    out = ['<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
           f'<w:comments xmlns:w="{W_NS}">']
    for cid, author, date, text in comments:
        initials = "".join(w[0] for w in re.findall(r"[A-Za-z]+", author))[:3].upper() or "RSI"
        out.append(f'<w:comment w:id="{cid}" w:author="{escape(author)}" '
                   f'w:date="{date}" w:initials="{initials}">')
        for i, line in enumerate(text.split("\n")):
            out.append("<w:p><w:r><w:t xml:space=\"preserve\">"
                       f"{escape(line)}</w:t></w:r></w:p>")
        out.append("</w:comment>")
    out.append("</w:comments>")
    return "".join(out)


def _inject_comments(path: Path, comments):
    if not comments:
        return
    path = Path(path)
    data = {}
    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        for n in names:
            data[n] = z.read(n)
    # 1) content types
    ct = data["[Content_Types].xml"].decode("utf-8")
    if "comments+xml" not in ct:
        ct = ct.replace("</Types>",
                        f'<Override PartName="/word/comments.xml" ContentType="{CT_COMMENTS}"/></Types>')
        data["[Content_Types].xml"] = ct.encode("utf-8")
    # 2) document rels
    rels_name = "word/_rels/document.xml.rels"
    rels = data[rels_name].decode("utf-8")
    if REL_COMMENTS not in rels:
        rid = "rIdComments100"
        rels = rels.replace("</Relationships>",
                            f'<Relationship Id="{rid}" Type="{REL_COMMENTS}" '
                            f'Target="comments.xml"/></Relationships>')
        data[rels_name] = rels.encode("utf-8")
    # 3) the part itself
    data["word/comments.xml"] = _comments_xml(comments).encode("utf-8")
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        for n, b in data.items():
            z.writestr(n, b)


# ---------------------------------------------------------------------------
# orchestration
# ---------------------------------------------------------------------------
def _recompute_ctx(manifest: dict, slots: dict) -> dict:
    """Mirror assemble's ctx build (flags + context_defaults + pins) so omitted
    blocks can be rendered with the matter's real facts."""
    s = dict(slots); s.pop("_meta", None)
    ctx = A.derive_flags(s)
    if manifest.get("context_defaults"):
        ctx = {**manifest["context_defaults"], **ctx}
    for slot, spec in (manifest.get("discriminators") or {}).items():
        if spec.get("pin") is not None:
            ctx[slot] = spec["pin"]
    return ctx


def export(manifest_path, slots: dict, out, author=DEF_AUTHOR, date=DEF_DATE):
    filled, base, manifest, blocks = build_records(manifest_path, slots)
    env = R.make_env()
    ctx = _recompute_ctx(manifest, slots)
    facts = filled.get("facts") or {}
    ids = _Ids()
    comments: list = []

    doc = Document()
    X._set_default_font(doc)
    n = doc.styles["Normal"]
    n.font.name = "Times New Roman"; n.font.size = Pt(12)
    _header(doc, filled, author)

    # index base decisions by block_id (for matching included blocks)
    base_by_id = {}
    if base:
        for d in base.get("decisions", []):
            if d.get("included"):
                base_by_id.setdefault(d.get("block_id"), []).append(d)

    if filled.get("title"):
        bt = (base or {}).get("title") if base else None
        emit_paragraph(doc, "title", bt or filled["title"], filled["title"],
                       facts, None, ids, comments, author, date)

    used = {}
    for d in filled.get("decisions", []):
        if not d.get("included"):
            # only show CONDITIONAL omissions (a real when-gate that fired false)
            if d.get("when") is not None and d.get("when_result") is False:
                emit_omission(doc, d, ctx, env, blocks, ids, comments, author, date)
            continue
        bid = d.get("block_id")
        authority = d.get("statutory_citation")
        # find the aligned base decision (same block_id, same occurrence)
        k = used.get(bid, 0); used[bid] = k + 1
        bd = None
        if bid in base_by_id and k < len(base_by_id[bid]):
            bd = base_by_id[bid][k]
        filled_text = d.get("rendered", "")
        base_text = (bd or {}).get("rendered", filled_text)
        role = (d.get("structure") or {}).get("role") or \
               ("chrome" if d.get("kind") == "chrome" else "body")
        if role not in S.ROLES:
            role = "body"
        # grids: diff the flat aligned-text view line by line
        if d.get("grid") is not None:
            ft_lines = S.grid_to_text(d["grid"]).split("\n")
            bt_lines = S.grid_to_text(bd["grid"]).split("\n") if (bd and bd.get("grid")) else ft_lines
            for i, fl in enumerate(ft_lines):
                bl = bt_lines[i] if i < len(bt_lines) else ""
                emit_paragraph(doc, "chrome", bl, fl, facts, authority, ids,
                               comments, author, date)
            continue
        fparas = S._split_paragraphs(filled_text) or [filled_text]
        bparas = S._split_paragraphs(base_text) or [base_text]
        if len(fparas) != len(bparas):
            # paragraph counts diverged (e.g. a within-block conditional) — diff
            # the whole block as one unit rather than mis-pairing paragraphs
            emit_paragraph(doc, role, base_text, filled_text, facts, authority,
                           ids, comments, author, date)
            continue
        for bp, fp in zip(bparas, fparas):
            emit_paragraph(doc, role, bp, fp, facts, authority, ids, comments,
                           author, date)

    Path(out).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    _inject_comments(out, comments)
    return out, {"inputs": ids.comment, "changes": ids.change,
                 "comments": len(comments),
                 "base_ok": base is not None}


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--pack", help="template pack directory "
                    "(default: $TEMPLATE_FORGE_PACK or the bundled example_pack)")
    ap.add_argument("--manifest", required=True)
    ap.add_argument("--slots", required=True)
    ap.add_argument("--out", required=True)
    ap.add_argument("--author", default=DEF_AUTHOR)
    ap.add_argument("--date", default=DEF_DATE)
    a = ap.parse_args()
    if a.pack:
        A.paths.set_pack(a.pack)
    out, stats = export(Path(a.manifest), json.load(open(a.slots)), a.out,
                        author=a.author, date=a.date)
    print(f"wrote {out}  ({stats['comments']} comments, {stats['changes']} "
          f"tracked changes, base_ok={stats['base_ok']})")


if __name__ == "__main__":
    main()
