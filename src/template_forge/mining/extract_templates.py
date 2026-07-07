#!/usr/bin/env python3
"""
DOCX Template Extractor - Extract and cluster legal document templates
Usage:
  extract_templates.py --phase inventory [--limit N]
  extract_templates.py --phase cluster
  extract_templates.py --phase templates
"""

import argparse
import json
import logging
import re
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import yaml
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

logger = logging.getLogger(__name__)


class DocxTemplateExtractor:
    """Extract and cluster DOCX templates from a document corpus.

    Fully config-driven: pass ``source_dir`` (where YOUR .docx documents live)
    and ``out_dir`` (where extracted templates/schemas are written). No paths are
    hardcoded — point it at your own corpus.
    """

    def __init__(self, source_dir: str, out_dir: str,
                 config: dict | None = None,
                 similarity_threshold: float = 0.85):
        """Initialize extractor.

        Args:
            source_dir: directory to walk for ``*.docx`` inputs.
            out_dir: directory for inventory/clusters/templates/schemas outputs.
            config: optional settings dict (reserved for future options).
            similarity_threshold: TF-IDF cosine threshold for clustering.
        """
        self.config = config or {}
        self.source_dir = Path(source_dir)
        self.out_dir = Path(out_dir)
        self.similarity_threshold = similarity_threshold

        # Create output directories
        (self.out_dir / 'templates').mkdir(parents=True, exist_ok=True)
        (self.out_dir / 'schemas').mkdir(parents=True, exist_ok=True)

        # Configure a file log handler under out_dir (idempotent).
        if not any(isinstance(h, logging.FileHandler) for h in logger.handlers):
            logger.setLevel(logging.INFO)
            fmt = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
            fh = logging.FileHandler(self.out_dir / 'extraction.log')
            fh.setFormatter(fmt)
            sh = logging.StreamHandler()
            sh.setFormatter(fmt)
            logger.addHandler(fh)
            logger.addHandler(sh)

        # Checkpoint file for resume capability
        self.checkpoint_file = str(self.out_dir / 'checkpoint.json')

        # Document type patterns
        self.doc_type_patterns = {
            'contract': r'\b(contract|agreement)\b',
            'letter': r'\b(dear|sincerely|regards)\b',
            'motion': r'\b(motion|moves|memorandum of law)\b',
            'deed': r'\b(deed|grantor|grantee)\b',
            'lease': r'\b(lease|lessor|lessee|tenant|landlord)\b',
            'will': r'\b(will|testament|testator|bequest)\b',
            'pleading': r'\b(plaintiff|defendant|complaint|answer)\b',
            'memorandum': r'\b(memorandum|memo)\b'
        }

        # Entity extraction patterns
        self.entity_patterns = {
            'date': r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}\b',
            'amount': r'\$[\d,]+(?:\.\d{2})?|\b\d+(?:,\d{3})*(?:\.\d{2})?\s*(?:dollars|USD)\b',
            'phone': r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b|\(\d{3}\)\s*\d{3}[-.]?\d{4}\b',
            'case_number': r'\b(?:Case|Docket)\s+(?:No\.?|Number|#)\s*[A-Z0-9-]+\b',
            'file_number': r'\b(?:File|Matter)\s+(?:No\.?|Number|#)\s*[A-Z0-9-]+\b',
            'address': r'\b\d+\s+[A-Za-z0-9\s,]+(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Lane|Ln|Drive|Dr|Court|Ct|Way|Place|Pl)\b',
            'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        }

    def phase1_inventory(self, limit: Optional[int] = None):
        """Phase 1: Walk directory, extract text, and inventory documents"""
        logger.info("=== PHASE 1: INVENTORY & EXTRACT ===")

        # Load checkpoint if exists
        checkpoint = self._load_checkpoint()
        processed_files = set(checkpoint.get('processed_files', []))

        # Find all .docx files
        source_dir = self.source_dir
        docx_files = [
            f for f in source_dir.rglob('*.docx')
            if not f.name.startswith('~$') and str(f) not in processed_files
        ]

        logger.info(f"Found {len(docx_files)} unprocessed .docx files")

        if limit:
            docx_files = docx_files[:limit]
            logger.info(f"Processing limited to {limit} files")

        inventory_file = self.out_dir / 'inventory.jsonl'

        processed_count = len(processed_files)
        error_count = 0

        with open(inventory_file, 'a') as f:
            for idx, docx_path in enumerate(docx_files, 1):
                try:
                    # Extract document info
                    doc_info = self._extract_document_info(docx_path)

                    # Write to JSONL
                    f.write(json.dumps(doc_info) + '\n')
                    f.flush()

                    processed_files.add(str(docx_path))
                    processed_count += 1

                    # Checkpoint every 100 files
                    if processed_count % 100 == 0:
                        self._save_checkpoint(list(processed_files))
                        logger.info(f"Processed {processed_count} files (checkpoint saved)")

                    if idx % 10 == 0:
                        logger.info(f"Progress: {idx}/{len(docx_files)} files")

                except Exception as e:
                    error_count += 1
                    logger.error(f"Error processing {docx_path}: {e}")
                    continue

        # Final checkpoint
        self._save_checkpoint(list(processed_files))

        logger.info(f"✓ Phase 1 complete: {processed_count} processed, {error_count} errors")
        logger.info(f"Inventory saved to: {inventory_file}")

    def _extract_document_info(self, docx_path: Path) -> Dict:
        """Extract text and metadata from a DOCX file"""
        doc = Document(docx_path)

        # Extract text
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        full_text = '\n'.join(paragraphs)

        # Count words
        word_count = len(full_text.split())

        # Detect document type
        doc_type = self._detect_document_type(full_text)

        # Get file size
        file_size = docx_path.stat().st_size

        return {
            'path': str(docx_path),
            'filename': docx_path.name,
            'file_size': file_size,
            'paragraph_count': len(paragraphs),
            'word_count': word_count,
            'doc_type': doc_type,
            'text': full_text[:10000],  # Limit text to first 10k chars for storage
            'text_hash': hash(full_text) % (2**32),  # Simple hash for dedup
            'extracted_at': datetime.now().isoformat()
        }

    def _detect_document_type(self, text: str) -> str:
        """Detect document type based on keyword patterns"""
        text_lower = text.lower()

        # Score each type
        scores = {}
        for doc_type, pattern in self.doc_type_patterns.items():
            matches = len(re.findall(pattern, text_lower, re.IGNORECASE))
            scores[doc_type] = matches

        # Return type with highest score, or 'unknown'
        if max(scores.values()) > 0:
            return max(scores, key=scores.get)
        return 'unknown'

    def phase2_cluster(self):
        """Phase 2: Cluster similar documents using TF-IDF"""
        logger.info("=== PHASE 2: SIMILARITY CLUSTERING ===")

        # Load inventory
        inventory_file = self.out_dir / 'inventory.jsonl'
        if not inventory_file.exists():
            logger.error("Inventory file not found. Run phase 1 first.")
            return

        documents = []
        with open(inventory_file, 'r') as f:
            for line in f:
                documents.append(json.loads(line))

        logger.info(f"Loaded {len(documents)} documents from inventory")

        # Group by document type first
        by_type = defaultdict(list)
        for doc in documents:
            by_type[doc['doc_type']].append(doc)

        all_clusters = []

        for doc_type, docs in by_type.items():
            logger.info(f"Clustering {len(docs)} documents of type '{doc_type}'")

            if len(docs) < 2:
                # Single document - create single cluster
                all_clusters.append({
                    'cluster_id': len(all_clusters),
                    'doc_type': doc_type,
                    'size': 1,
                    'representative': docs[0]['path'],
                    'members': [docs[0]['path']]
                })
                continue

            # Extract texts
            texts = [d['text'] for d in docs]

            # TF-IDF vectorization
            vectorizer = TfidfVectorizer(
                max_features=5000,
                stop_words='english',
                ngram_range=(1, 2)
            )

            try:
                tfidf_matrix = vectorizer.fit_transform(texts)
            except ValueError:
                logger.warning(f"Could not vectorize {doc_type} documents, skipping clustering")
                continue

            # Calculate similarity matrix
            similarity_matrix = cosine_similarity(tfidf_matrix)

            # Cluster using threshold
            threshold = self.similarity_threshold
            visited = set()
            clusters = []

            for i in range(len(docs)):
                if i in visited:
                    continue

                # Find all similar documents
                cluster_members = [i]
                for j in range(i + 1, len(docs)):
                    if j not in visited and similarity_matrix[i][j] > threshold:
                        cluster_members.append(j)
                        visited.add(j)

                visited.add(i)

                # Pick representative (longest document)
                rep_idx = max(cluster_members, key=lambda idx: docs[idx]['word_count'])

                clusters.append({
                    'cluster_id': len(all_clusters) + len(clusters),
                    'doc_type': doc_type,
                    'size': len(cluster_members),
                    'representative': docs[rep_idx]['path'],
                    'members': [docs[idx]['path'] for idx in cluster_members]
                })

            all_clusters.extend(clusters)
            logger.info(f"  → Created {len(clusters)} clusters for {doc_type}")

        # Save clusters
        clusters_file = self.out_dir / 'clusters.json'
        with open(clusters_file, 'w') as f:
            json.dump(all_clusters, f, indent=2)

        logger.info(f"✓ Phase 2 complete: {len(all_clusters)} clusters created")
        logger.info(f"Clusters saved to: {clusters_file}")

    def phase3_templates(self):
        """Phase 3: Generate Jinja2 templates from cluster representatives"""
        logger.info("=== PHASE 3: TEMPLATE GENERATION ===")

        # Load clusters
        clusters_file = self.out_dir / 'clusters.json'
        if not clusters_file.exists():
            logger.error("Clusters file not found. Run phase 2 first.")
            return

        with open(clusters_file, 'r') as f:
            clusters = json.load(f)

        logger.info(f"Generating templates for {len(clusters)} clusters")

        templates_created = 0

        for cluster in clusters:
            try:
                # Load representative document
                doc_path = Path(cluster['representative'])
                doc = Document(doc_path)

                # Extract full text
                paragraphs = [p.text for p in doc.paragraphs]
                full_text = '\n'.join(paragraphs)

                # Generate template
                template_text, schema = self._generate_template(full_text, cluster)

                # Save template
                template_filename = f"template_{cluster['cluster_id']:04d}_{cluster['doc_type']}.j2"
                template_path = self.out_dir / 'templates' / template_filename

                with open(template_path, 'w') as f:
                    f.write(template_text)

                # Save schema
                schema_filename = f"schema_{cluster['cluster_id']:04d}_{cluster['doc_type']}.yaml"
                schema_path = self.out_dir / 'schemas' / schema_filename

                with open(schema_path, 'w') as f:
                    yaml.dump(schema, f, default_flow_style=False)

                templates_created += 1

                if templates_created % 10 == 0:
                    logger.info(f"Generated {templates_created} templates...")

            except Exception as e:
                logger.error(f"Error generating template for cluster {cluster['cluster_id']}: {e}")
                continue

        logger.info(f"✓ Phase 3 complete: {templates_created} templates created")

    def _generate_template(self, text: str, cluster: Dict) -> Tuple[str, Dict]:
        """Convert document text to Jinja2 template with variable extraction"""

        template_text = text
        variables = {}
        var_counter = defaultdict(int)

        # Replace entities with Jinja2 variables
        for entity_type, pattern in self.entity_patterns.items():
            matches = re.finditer(pattern, template_text, re.IGNORECASE)

            # Collect matches (reverse order to maintain indices)
            replacements = []
            for match in matches:
                replacements.append((match.start(), match.end(), match.group()))

            # Replace from end to start
            for start, end, value in reversed(replacements):
                var_counter[entity_type] += 1
                var_name = f"{entity_type}_{var_counter[entity_type]}"

                # Store variable info
                variables[var_name] = {
                    'type': entity_type,
                    'example': value,
                    'required': True
                }

                # Replace in text
                template_text = (
                    template_text[:start] +
                    '{{ ' + var_name + ' }}' +
                    template_text[end:]
                )

        # Create schema
        schema = {
            'template_id': cluster['cluster_id'],
            'doc_type': cluster['doc_type'],
            'cluster_size': cluster['size'],
            'variables': variables,
            'created_at': datetime.now().isoformat()
        }

        return template_text, schema

    def _load_checkpoint(self) -> Dict:
        """Load processing checkpoint"""
        if Path(self.checkpoint_file).exists():
            with open(self.checkpoint_file, 'r') as f:
                return json.load(f)
        return {}

    def _save_checkpoint(self, processed_files: List[str]):
        """Save processing checkpoint"""
        with open(self.checkpoint_file, 'w') as f:
            json.dump({
                'processed_files': processed_files,
                'updated_at': datetime.now().isoformat()
            }, f)


def main():
    parser = argparse.ArgumentParser(description='DOCX Template Extractor')
    parser.add_argument('--phase', required=True,
                       choices=['inventory', 'cluster', 'templates'],
                       help='Processing phase to run')
    parser.add_argument('--source-dir', required=True,
                       help='Directory of your .docx documents to mine')
    parser.add_argument('--out-dir', required=True,
                       help='Output directory for inventory/clusters/templates/schemas')
    parser.add_argument('--threshold', type=float, default=0.85,
                       help='TF-IDF cosine similarity clustering threshold')
    parser.add_argument('--limit', type=int,
                       help='Limit number of files to process (inventory phase only)')

    args = parser.parse_args()

    extractor = DocxTemplateExtractor(
        source_dir=args.source_dir, out_dir=args.out_dir,
        similarity_threshold=args.threshold)

    if args.phase == 'inventory':
        extractor.phase1_inventory(limit=args.limit)
    elif args.phase == 'cluster':
        extractor.phase2_cluster()
    elif args.phase == 'templates':
        extractor.phase3_templates()


if __name__ == '__main__':
    main()
